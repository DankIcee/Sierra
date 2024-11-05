from flask import Flask, render_template, request, send_file
import os
from dotenv import load_dotenv
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_TAB_ALIGNMENT
from PyPDF2 import PdfReader
from anthropic import Anthropic
from PIL import Image, ExifTags
import io
import fitz  # PyMuPDF for image extraction
import re
import logging
import cv2
import numpy as np

app = Flask(__name__)
app.config['MAX_CONTENT_LENGTH'] = 30 * 1024 * 1024  # 30MB max file size
app.config['SEND_FILE_MAX_AGE_DEFAULT'] = 0
app.config['PERMANENT_SESSION_LIFETIME'] = 600  # 10 minutes

logging.basicConfig(level=logging.DEBUG, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

load_dotenv()
anthropic_api_key = os.getenv('ANTHROPIC_API_KEY')
client = Anthropic(api_key=anthropic_api_key)
modern_font = "Calibri"

def detect_face(image):
    try:
        opencv_img = cv2.cvtColor(np.array(image), cv2.COLOR_RGB2BGR)
        face_cascade = cv2.CascadeClassifier(cv2.data.haarcascades + 'haarcascade_frontalface_default.xml')
        gray = cv2.cvtColor(opencv_img, cv2.COLOR_BGR2GRAY)
        faces = face_cascade.detectMultiScale(
            gray,
            scaleFactor=1.1,
            minNeighbors=5,
            minSize=(30, 30)
        )
        return len(faces) > 0
    except Exception as e:
        logger.error(f"Error in face detection: {str(e)}")
        return False

def process_extracted_image(image, min_face_size=(30, 30)):
    try:
        if image is None:
            return None

        opencv_img = cv2.cvtColor(np.array(image), cv2.COLOR_RGB2BGR)
        face_cascade = cv2.CascadeClassifier(cv2.data.haarcascades + 'haarcascade_frontalface_default.xml')
        gray = cv2.cvtColor(opencv_img, cv2.COLOR_BGR2GRAY)
        faces = face_cascade.detectMultiScale(
            gray,
            scaleFactor=1.1,
            minNeighbors=5,
            minSize=min_face_size
        )
        
        if len(faces) == 0:
            rotated = image.rotate(90, expand=True)
            opencv_rotated = cv2.cvtColor(np.array(rotated), cv2.COLOR_RGB2BGR)
            gray_rotated = cv2.cvtColor(opencv_rotated, cv2.COLOR_BGR2GRAY)
            faces_rotated = face_cascade.detectMultiScale(
                gray_rotated,
                scaleFactor=1.1,
                minNeighbors=5,
                minSize=min_face_size
            )
            if len(faces_rotated) > 0:
                return rotated
            return None
        return image
    except Exception as e:
        logger.error(f"Error processing image: {str(e)}")
        return None

def fix_image_orientation(image):
    try:
        original_image = image.copy()
        rotations = [0, 90, 180, 270]
        best_rotation = 0
        face_found = False

        try:
            for orientation in ExifTags.TAGS.keys():
                if ExifTags.TAGS[orientation] == 'Orientation':
                    break
            exif = dict(image._getexif().items())
            if exif[orientation] == 3:
                image = image.rotate(180, expand=True)
            elif exif[orientation] == 6:
                image = image.rotate(270, expand=True)
            elif exif[orientation] == 8:
                image = image.rotate(90, expand=True)
            
            if detect_face(image):
                return image
        except (AttributeError, KeyError, IndexError):
            pass

        for angle in rotations:
            rotated = original_image.rotate(angle, expand=True)
            if detect_face(rotated):
                face_found = True
                best_rotation = angle
                break

        if face_found:
            return original_image.rotate(best_rotation, expand=True)

        width, height = original_image.size
        if width > height and width / height > 1.5:
            return original_image.rotate(90, expand=True)

        return original_image
    except Exception as e:
        logger.error(f"Error in image orientation: {str(e)}")
        return image

def extract_cv_information(cv_text):
    try:
        logger.debug("Processing CV text")
        
        prompt = (
            "IMPORTANT: The following CV text may be quite long. Please ensure you allocate sufficient memory "
            "to process the entire content thoroughly.\n\n"
            "You are an AI assistant that extracts and formats information from resumes. "
            "Given the following CV text, please extract and format ALL relevant information "
            "into a well-structured, professional-looking resume. Only use the information provided in the CV text below. "
            "Do not add any information that is not present in the given CV. Do not add in you're generic prompt of 'Here's your formatted CV' and etc. Follow these strict guidelines:\n\n"
            "1. CRITICAL - Content Preservation:\n"
            "   - Keep 98% of content exactly as written\n"
            "   - Only fix basic grammar and punctuation\n"
            "   - Do not paraphrase or rewrite content\n"
            "   - Keep all technical terms exactly as written\n"
            "   - Preserve all metrics and numbers exactly\n"
            "   - Maintain all project descriptions exactly\n"
            "2. Grammar Fixing Rules:\n"
            "   - Fix only obvious spelling mistakes\n"
            "   - Add missing periods at end of sentences\n"
            "   - Fix basic capitalization only where clearly wrong\n"
            "   - Fix only obvious punctuation errors\n"
            "   - Do not rephrase or rewrite sentences\n"
            "3. Preserve the structure of the CV exactly as provided. Extract ALL sections in their original order.\n"
            "4. Extract ALL information from the CV, including full sentences, paragraphs, and bullet points exactly as written.\n"
            "5. Use the exact words and phrases from the original CV.\n"
            "6. Do not add any information that is not present in the original CV.\n"
            "7. Maintain all specific details, metrics, numbers, and achievements exactly as written.\n"
            "8. Handle educational qualifications thoroughly:\n"
            "   - Extract ALL degrees, certifications, and qualifications exactly as written\n"
            "   - Include all certification bodies, institutions, and authorizing organizations\n"
            "   - Preserve full names of certifying bodies (e.g., 'KHDA, Ministry of Education')\n"
            "   - Keep complete qualification names (e.g., 'Certified Human Resource Management Professional (CHRMP)')\n"
            "   - Maintain all location information (e.g., 'Dubai, UAE')\n"
            "   - Keep certification years/dates exactly as shown\n"
            "9. Use the following format (BE VERY THOROUGH) for structuring the extracted information:\n"
            "[NAME]Full Name\n"   
            "[SECTION]Section Heading\n"
            "[COMPANY]Company Name, Location\n"
            "[JOBTITLE]Job Title | Date Range\n"
            "[BULLET]Bullet point\n"
            "[EDUCATION]Degree/Qualification Name | Date Range\n"
            "[INSTITUTION]Institution Name, Location\n"
            "[CERTBODY]Certifying Body/Authority Details\n"
            "[CERTIFICATION]Certification Name | Date\n"
            "[CERTORG]Organization Name, Location\n"
            "[NORMAL]Normal text\n"
            "[SUBHEADING]Subheading or Category Title\n\n"
            "10. Ensure that dates are consistently formatted as 'Month Year - Month Year' or 'Month Year - Present' or just 'Year' as appropriate.\n"
            "11. Include countries or cities where the person has worked or studied, ensuring they are properly formatted with correct capitalization (e.g., 'Dubai, United Arab Emirates'). Do not include any asterisks or special characters.\n"
            "12. Remove all contact-based information if possible, also, parent names, passport number, marital status, religion or similar.\n"
            "13. CRITICAL: For the Summary section:\n"
            "    - Look for and include ALL profile highlights, key skills, or professional summary information\n"
            "    - Search for sections labeled as 'Profile', 'Summary', 'Professional Summary', 'Profile Highlights', 'Key Skills', etc.\n"
            "    - Format the Summary section with clear structure:\n"
            "      * Start with a brief overview paragraph if available\n"
            "      * Use [SUBHEADING] for category titles (e.g., 'Core Competencies', 'Technical Skills')\n"
            "      * Follow each subheading with relevant [BULLET] points\n"
            "      * Maintain original grouping of skills and highlights\n"
            "14. Organize the content in the following order (if available):\n"
            "    - Summary/Profile information/Profile skills/Profile Highlights/Profile (Name it as 'Summary')(If no title is provided, recognize it and then place it in the Summary section)\n"
            "    - Work Experience and etc (Name as 'Experience')\n"
            "    - Education and Qualifications (All degrees, certifications, and qualifications)\n"
            "    - Technical Skills (if exists)\n"
            "    - Language Proficiency (if mentioned)\n"
            "    - Professional Training (if exists)\n"
            "    - Awards & Achievements (if exists)\n"
            "15. For any section not listed above, create an appropriate professional section name with approriate bulleting for each.\n"
            "16. DO NOT use tabs to separate information. Use the '|' character to separate titles and dates.\n"
            "17. Ensure education entries and certification entries have dates on the same line, separated by ' | '.\n"
            "18. When processing educational qualifications and certifications, preserve ALL details including:\n"
            "    - Full qualification names with any abbreviations\n"
            "    - Complete certification body names\n"
            "    - All locations and institutions\n"
            "    - All dates and year information\n"
            "19. ALWAYS place skills in the lower half of the CV.\n"
            "20. For Professional Training section:\n"
            "    - Format all training items as bullet points using [BULLET]\n"
            "    - Include dates and certification details within the bullet points\n"
            "    - Maintain chronological order if dates are provided\n"
            "21. CRITICAL: Grammar and Formatting Rules:\n"
            "    - Make only essential grammar fixes that don't change meaning\n"
            "    - Fix only obvious spelling mistakes\n"
            "    - Add missing punctuation only where clearly needed\n"
            "    - Keep all technical terms exactly as written\n"
            "    - Maintain all abbreviations as shown\n"
            "Please provide the formatted CV content, ready to be inserted into a Word document. "
            "Remember to maintain the original content as much as possible except information not relevant in a CV."
            "Do not create a summary if the person has not provided any."
        )

        response = client.messages.create(
            model="claude-3-5-sonnet-20240620",
            max_tokens=8192,
            temperature=0.2,
            messages=[{"role": "user", "content": prompt + f"\n\nCV Text:\n{cv_text}"}]
        )
        
        logger.debug(f"API Response type: {type(response.content)}")
        
        if isinstance(response.content, list) and len(response.content) > 0:
            if hasattr(response.content[0], 'text'):
                formatted_cv = response.content[0].text
            else:
                formatted_cv = str(response.content[0])
        else:
            formatted_cv = str(response.content)
        
        logger.debug(f"Processed CV: {formatted_cv[:1000]}...")
        return formatted_cv
            
    except Exception as e:
        logger.error(f"Error processing CV: {str(e)}")
        raise

def extract_text_from_pdf(pdf_path):
    try:
        with open(pdf_path, 'rb') as file:
            reader = PdfReader(file)
            text = ''
            for page in reader.pages:
                text += page.extract_text() or ''
        return text
    except Exception as e:
        logger.error(f"Error extracting text from PDF: {str(e)}")
        raise

def extract_image_from_pdf(pdf_path):
    try:
        doc = fitz.open(pdf_path)
        images = []
        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            image_list = page.get_images(full=True)
            for img_index, img in enumerate(image_list):
                xref = img[0]
                base_image = doc.extract_image(xref)
                image_bytes = base_image["image"]
                image_stream = io.BytesIO(image_bytes)
                image = Image.open(image_stream)
                processed_image = process_extracted_image(image)
                if processed_image:
                    images.append((processed_image, processed_image.size[0] * processed_image.size[1]))
        
        if images:
            largest_image = max(images, key=lambda x: x[1])[0]
            return largest_image
        return None
    except Exception as e:
        logger.error(f"Error extracting image from PDF: {str(e)}")
        return None

def extract_text_from_docx(docx_path):
    try:
        doc = Document(docx_path)
        full_text = []

        for para in doc.paragraphs:
            full_text.append(para.text)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text.append(cell.text)

        return '\n'.join(full_text)
    except Exception as e:
        logger.error(f"Error extracting text from DOCX: {str(e)}")
        raise

def extract_image_from_docx(docx_path):
    try:
        doc = Document(docx_path)
        images = []
        for rel in doc.part.rels:
            if "image" in doc.part.rels[rel].target_ref:
                image = doc.part.rels[rel].target_part.blob
                image_stream = io.BytesIO(image)
                img = Image.open(image_stream)
                processed_image = process_extracted_image(img)
                if processed_image:
                    images.append((processed_image, processed_image.size[0] * processed_image.size[1]))
        
        if images:
            largest_image = max(images, key=lambda x: x[1])[0]
            return largest_image
        return None
    except Exception as e:
        logger.error(f"Error extracting image from DOCX: {str(e)}")
        return None

def clean_and_normalize_text(text):
    text = ' '.join(text.split())
    text = re.sub(r'\n+', '\n', text)
    text = re.sub(r'[^\x20-\x7E\n]', '', text)
    return text

def clean_text(text):
    if isinstance(text, list):
        text = ' '.join(text)
    text = text.replace("[NORMAL]", "").replace("[TAB]", "\t")
    text = text.replace("*", "")
    return text

def create_word_doc(output_path, formatted_cv, cv_image=None):
    try:
        doc = Document('templates/naas_template.docx')
        styles = doc.styles
        
        cv_normal_style = styles['Normal']
        cv_normal_style.font.size = Pt(11)
        cv_normal_style.font.name = modern_font
        cv_normal_style.font.color.rgb = RGBColor(0, 0, 0)
        cv_normal_style.paragraph_format.space_after = Pt(1)
        
        style_definitions = {
            'CV_ApplicantName': {
                'size': 18,
                'bold': True,
                'alignment': WD_PARAGRAPH_ALIGNMENT.CENTER,
                'space_after': 3
            },
            'CV_SectionHeading': {
                'size': 14,
                'bold': True,
                'space_after': 3
            },
            'CV_CompanyName': {
                'size': 12,
                'bold': True,
                'space_after': 1
            },
            'CV_JobTitle': {
                'size': 11,
                'italic': True,
                'space_after': 1
            },
            'CV_Education': {
                'size': 11,
                'bold': True,
                'space_after': 1
            },
            'CV_Institution': {
                'size': 11,
                'italic': True,
                'space_after': 1
            },
            'CV_Certification': {
                'size': 11,
                'bold': True,
                'space_after': 1
            },
            'CV_CertBody': {
                'size': 11,
                'italic': True,
                'space_after': 0
            },
            'CV_Subheading': {
                'size': 11,
                'italic': True,
                'space_after': 1
            }
        }

        for style_name, properties in style_definitions.items():
            style = styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
            style.font.size = Pt(properties['size'])
            style.font.name = modern_font
            style.font.color.rgb = RGBColor(0, 0, 0)
            if properties.get('bold'):
                style.font.bold = True
            if properties.get('italic'):
                style.font.italic = True
            if properties.get('alignment'):
                style.paragraph_format.alignment = properties['alignment']
            style.paragraph_format.space_after = Pt(properties['space_after'])
            style.paragraph_format.line_spacing = 1.0

        cleaned_cv = clean_text(formatted_cv)
        lines = cleaned_cv.split('\n')
        applicant_name = ""
        current_section = ""
        in_experience_section = False
        in_education_section = False
        in_certification_section = False
        previous_line_type = None

        if cv_image:
            image_stream = io.BytesIO()
            cv_image.save(image_stream, format="PNG")
            image_stream.seek(0)
            doc.add_picture(image_stream, width=Inches(1.5))
            doc.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        for line in lines:
            line = line.strip()
            if not line:
                continue

            if line.startswith('[NAME]'):
                p = doc.add_paragraph(line[len('[NAME]'):].strip(), style='CV_ApplicantName')
                applicant_name = line[len('[NAME]'):].strip()

            elif line.startswith('[SECTION]'):
                section_name = line[len('[SECTION]'):].strip().upper()
                current_section = section_name.lower()
                in_experience_section = (current_section == 'experience')
                in_education_section = (current_section == 'education')
                in_certification_section = (current_section == 'professional certifications' or 
                                         current_section == 'certifications' or 
                                         current_section == 'certification')
                
                if previous_line_type:
                    doc.add_paragraph()
                
                p = doc.add_paragraph(section_name, style='CV_SectionHeading')
                previous_line_type = 'section'

            elif line.startswith('[SUBHEADING]'):
                p = doc.add_paragraph(line[len('[SUBHEADING]'):].strip(), style='CV_Subheading')
                previous_line_type = 'subheading'

            elif line.startswith('[COMPANY]'):
                if previous_line_type and previous_line_type != 'section':
                    doc.add_paragraph()
                
                p = doc.add_paragraph(style='CV_CompanyName')
                company_info = line[len('[COMPANY]'):].strip()
                p.add_run(company_info)
                previous_line_type = 'company'

            elif line.startswith('[JOBTITLE]'):
                p = doc.add_paragraph()
                job_info = line[len('[JOBTITLE]'):].strip()
                
                run = p.add_run(job_info)
                run.font.name = modern_font
                run.font.size = Pt(11)
                run.font.italic = True
                run.font.color.rgb = RGBColor(0, 0, 0)
                
                p.paragraph_format.space_after = Pt(1)
                p.paragraph_format.line_spacing = 1.0
                
                previous_line_type = 'jobtitle'

            elif line.startswith('[BULLET]'):
                p = doc.add_paragraph(style='List Bullet')
                p.text = line[len('[BULLET]'):].strip()
                p.paragraph_format.left_indent = Inches(0.25)
                if not (in_education_section or in_certification_section):
                    p.paragraph_format.space_after = Pt(1)
                else:
                    p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.space_before = Pt(0)
                previous_line_type = 'bullet'

            elif line.startswith('[EDUCATION]'):
                if previous_line_type and previous_line_type != 'section':
                    doc.add_paragraph()
                
                p = doc.add_paragraph(style='CV_Education')
                education_info = line[len('[EDUCATION]'):].strip()
                p.add_run(education_info)
                previous_line_type = 'education'

            elif line.startswith('[INSTITUTION]'):
                p = doc.add_paragraph(style='CV_Institution')
                institution_info = line[len('[INSTITUTION]'):].strip()
                p.add_run(institution_info)
                previous_line_type = 'institution'

            elif line.startswith('[CERTBODY]'):
                p = doc.add_paragraph(style='CV_CertBody')
                cert_body_info = line[len('[CERTBODY]'):].strip()
                p.add_run(cert_body_info)
                previous_line_type = 'certbody'

            elif line.startswith('[CERTIFICATION]'):
                if previous_line_type and previous_line_type != 'section':
                    doc.add_paragraph()
                
                p = doc.add_paragraph(style='CV_Certification')
                cert_info = line[len('[CERTIFICATION]'):].strip()
                p.add_run(cert_info)
                previous_line_type = 'certification'

            elif line.startswith('[CERTORG]'):
                p = doc.add_paragraph(style='CV_Institution')
                org_info = line[len('[CERTORG]'):].strip()
                p.add_run(org_info)
                previous_line_type = 'certorg'

            elif line:
                p = doc.add_paragraph(line, style='Normal')
                previous_line_type = 'normal'

        if not (in_education_section or in_certification_section):
            doc.add_paragraph()
            
        doc.save(output_path)
        return applicant_name
        
    except Exception as e:
        logger.error(f"Error creating Word document: {str(e)}")
        logger.error(f"Problematic CV content: {formatted_cv}")
        raise

@app.route('/health')
def health_check():
   return "OK", 200

@app.route('/')
def index():
   return render_template('upload.html')

@app.route('/upload', methods=['POST'])
def upload_file():
   if 'file' not in request.files:
       return "No file part", 400

   file = request.files['file']

   if file.filename == '':
       return "No selected file", 400

   if file and (file.filename.endswith('.pdf') or file.filename.endswith('.docx')):
       upload_folder = 'uploads'
       output_folder = 'outputs'

       os.makedirs(upload_folder, exist_ok=True)
       os.makedirs(output_folder, exist_ok=True)

       file_path = os.path.join(upload_folder, file.filename)
       file.save(file_path)

       try:
           if file.filename.endswith('.pdf'):
               cv_text = extract_text_from_pdf(file_path)
               cv_image = extract_image_from_pdf(file_path)
           elif file.filename.endswith('.docx'):
               cv_text = extract_text_from_docx(file_path)
               cv_image = extract_image_from_docx(file_path)

           cv_text = clean_and_normalize_text(cv_text)
           logger.debug(f"Extracted and cleaned CV text: {cv_text[:1000]}...")

           formatted_cv = extract_cv_information(cv_text)
           
           temp_output_path = os.path.join(output_folder, 'temp_CV.docx')
           applicant_name = create_word_doc(temp_output_path, formatted_cv, cv_image)

           if not applicant_name:
               applicant_name = "Unknown"

           base_filename = f'{applicant_name.replace(" ", "_")}_CV.docx'
           final_output_path = os.path.join(output_folder, base_filename)

           counter = 1
           while os.path.exists(final_output_path):
               base_filename = f'{applicant_name.replace(" ", "_")}_CV({counter}).docx'
               final_output_path = os.path.join(output_folder, base_filename)
               counter += 1

           os.rename(temp_output_path, final_output_path)

           return send_file(final_output_path, as_attachment=True)

       except Exception as e:
           logger.error(f"Error processing file: {str(e)}")
           return f"Error processing file: {str(e)}", 500

       finally:
           if os.path.exists(file_path):
               os.remove(file_path)

   return "Unsupported file format. Please upload a PDF or DOCX.", 400

if __name__ == '__main__':
   app.run(debug=True)
